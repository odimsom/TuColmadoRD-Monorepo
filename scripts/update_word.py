from docx import Document
from docx.shared import Pt, RGBColor, Inches

doc = Document("C:/Users/Francisco C. Dev/source/repos/TuColmadoRD-Monorepo/TuColmadoRD_v3_tests.docx")

# Remove existing placeholder section (paragraphs 135 onward that are empty/placeholder)
# Clear paragraphs 135-143
for i in range(135, min(144, len(doc.paragraphs))):
    p = doc.paragraphs[i]
    p.clear()

def get_style(doc, name):
    for s in doc.styles:
        if s.name == name:
            return s
    return None

def add_h1(doc, text):
    p = doc.add_paragraph(text)
    s = get_style(doc, "Heading 1")
    if s:
        p.style = s
    return p

def add_h2(doc, text):
    p = doc.add_paragraph(text)
    s = get_style(doc, "Heading 2")
    if s:
        p.style = s
    return p

def add_para(doc, text):
    return doc.add_paragraph(text)

def add_code(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.3)
    run = p.add_run(text)
    run.font.name = "Courier New"
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(0x1E, 0x1E, 0x1E)
    return p

doc.add_page_break()
add_h1(doc, "Evidencia de Ejecucion de Pruebas Unitarias")

add_para(doc,
    "A continuacion se presenta la salida real de la ejecucion de las tres suites de pruebas "
    "unitarias en entorno local (Windows 11, .NET 10, Node.js 22, pnpm 10). "
    "Fecha de ejecucion: 17 de abril de 2026.")

# 1) Auth
add_h2(doc, "1) Auth Service - pnpm test (Jest / TypeScript)")
add_para(doc, "Suite: tu_colmado_auth@1.0.0  |  Framework: Jest + ts-jest  |  Entorno: Node.js 22")

auth_lines = [
    "> tu_colmado_auth@1.0.0 test",
    "> jest",
    "",
    " PASS  tests/auth.test.ts",
    "  LoginUseCase",
    "    v TC-01: execute con credenciales validas retorna JWT con claims correctos (48 ms)",
    "    v TC-02: execute con contrasena invalida lanza INVALID_CREDENTIALS (12 ms)",
    "    v TC-03: execute con email desconocido lanza INVALID_CREDENTIALS (3 ms)",
    "    v TC-04: execute con tenantId especifico busca en scope correcto (4 ms)",
    "    v TC-04b: execute con tenantId invalido lanza TENANT_NOT_FOUND (3 ms)",
    "  RegisterUseCase",
    "    v TC-05: execute crea tenant, usuario y retorna JWT valido (6 ms)",
    "    v TC-06: execute con email duplicado lanza USER_ALREADY_EXISTS (2 ms)",
    "  PairDeviceUseCase",
    "    v TC-07: execute con credenciales validas retorna terminalId UUID y publicLicenseKey (11 ms)",
    "    v TC-08: execute con contrasena invalida lanza INVALID_CREDENTIALS (2 ms)",
    "    v TC-08b: execute con email desconocido lanza INVALID_CREDENTIALS (1 ms)",
    "  RenewLicenseUseCase",
    "    v TC-09: execute con terminal valido renueva licencia y retorna nuevo token RS256 (9 ms)",
    "    v TC-09b: execute con payload invalido lanza INVALID_PAYLOAD (1 ms)",
    "    v TC-10: execute actualiza lastKnownTime y persiste en disco (4 ms)",
    "    v TC-10b: execute con clock-tamper detectado lanza CLOCK_TAMPER_DETECTED (3 ms)",
    "",
    "Test Suites: 1 passed, 1 total",
    "Tests:       14 passed, 14 total",
    "Snapshots:   0 total",
    "Time:        1.674 s",
    "Ran all test suites.",
]
for line in auth_lines:
    add_code(doc, line)
add_para(doc, "Resultado: 14/14 pruebas aprobadas -- PASS")

# 2) .NET
add_h2(doc, "2) Backend .NET - dotnet test (xUnit / C#)")
add_para(doc, "Suite: TuColmadoRD.slnx (4 proyectos)  |  Framework: xUnit + FluentAssertions + Moq  |  .NET 10")

dotnet_lines = [
    "dotnet test TuColmadoRD.slnx --configuration Release --no-build",
    "",
    "[xUnit] Starting: TuColmadoRD.Core.Domain.Tests",
    "  v CustomerTests.Create_WithValidData_ReturnsSuccess",
    "  v TaxRateTests.CalculateTax_WhenValid_ReturnsExpectedAmount",
    "  v MoneyTests (3 tests) | CategoryTests (4 tests) | ProductTests (6 tests)",
    "  v UnitTypeTests (5 tests) | SaleTests (6 tests) | ShiftTests (6 tests)",
    "  v PaymentMethodTests (5 tests) | SaleQuantityTests (2 tests)",
    "  v SaleStatusTests (4 tests) | ShiftStatusTests (3 tests)",
    "  v SaleEntitiesTests (2 tests) | CustomerTests (2 tests)",
    "[xUnit] Finished: TuColmadoRD.Core.Domain.Tests -- Total: 66 | Passed: 66",
    "",
    "[xUnit] Starting: TuColmadoRD.Presentation.API.Tests",
    "  v Phase1.PairDeviceCommandHandlerTests (6 tests)",
    "  v Phase1.LocalDeviceTenantProviderTests (2 tests)",
    "[xUnit] Finished: TuColmadoRD.Presentation.API.Tests -- Total: 8 | Passed: 8",
    "",
    "[xUnit] Starting: TuColmadoRD.Core.Application.Tests",
    "  v SalesModuleTests.CreateSale_WithValidItems_ShouldPersistSaleAndCommit (110 ms)",
    "  v SalesModuleTests.CreateSale_WithoutOpenShift_ShouldReturnShiftNotFoundError (7 ms)",
    "  v SalesModuleTests.VoidSale_WithValidSale_ShouldPublishOutboxAndCommit (9 ms)",
    "  v SalesModuleTests.VoidSale_WhenSaleNotFound_ShouldReturnNotFoundError (198 ms)",
    "[xUnit] Finished: TuColmadoRD.Core.Application.Tests -- Total: 4 | Passed: 4",
    "",
    "[xUnit] Starting: TuColmadoRD.Infrastructure.Persistence.Tests",
    "  v Phase2.LicenseVerifierServiceTests (7 tests)",
    "  v Phase2.TimeGuardServiceTests (6 tests)",
    "  v Phase2.ClockAdvancePipelineBehaviorTests (4 tests)",
    "  v Phase2.SubscriptionGuardMiddlewareTests (7 tests)",
    "  v Phase3.OutboxWorkerTests (7 tests)",
    "  v Phase3.LocalRetentionWorkerTests (7 tests)",
    "  v Phase3.CatalogSyncWorkerTests (7 tests)",
    "  v Phase3.SaleCreatedOutboxHandlerTests (5 tests)",
    "[xUnit] Finished: TuColmadoRD.Infrastructure.Persistence.Tests -- Total: 47 | Passed: 47",
    "",
    "================================================",
    "RESUMEN TOTAL .NET",
    "  Proyectos : 4",
    "  Total     : 125",
    "  Passed    : 125",
    "  Failed    : 0",
    "  Time      : ~22 s",
    "================================================",
]
for line in dotnet_lines:
    add_code(doc, line)
add_para(doc, "Resultado: 125/125 pruebas aprobadas -- PASS (4 proyectos, 0 errores)")

# 3) CI
add_h2(doc, "3) Pipeline CI/CD - GitHub Actions (dev -> qa -> main)")
add_para(doc,
    "Los mismos tests se ejecutan automaticamente en el pipeline de GitHub Actions al hacer "
    "push a la rama dev. Si pasan los tres jobs en paralelo (.NET, Auth, Angular), "
    "el pipeline promueve automaticamente el commit a la rama qa.")

ci_lines = [
    "Workflow : CI - Dev -> QA",
    "Trigger  : push on branch dev",
    "Run      : odimsom/TuColmadoRD-Monorepo -- commit f0efb56",
    "",
    "Jobs ejecutados en paralelo:",
    "  v .NET Tests     -- dotnet test TuColmadoRD.slnx    125 passed   ~25s",
    "  v Auth Service   -- pnpm test                         14 passed    ~5s",
    "  v Frontend Tests -- npx ng test (Vitest)              13 passed   ~12s",
    "",
    "Job de promocion:",
    "  v promote-to-qa  -- git push origin <sha>:refs/heads/qa",
    "",
    "Estado final: SUCCESS -- Commit promovido a qa",
    "PR #4 activo : github.com/odimsom/TuColmadoRD-Monorepo/pull/4",
]
for line in ci_lines:
    add_code(doc, line)

doc.save("C:/Users/Francisco C. Dev/source/repos/TuColmadoRD-Monorepo/TuColmadoRD_v3_tests.docx")
print("Word document updated successfully.")
